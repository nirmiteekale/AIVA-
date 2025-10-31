from docx import Document

def markdown_to_docx(md_text: str, out_path: str):
    """
    Minimal Markdown -> DOCX converter for headings and paragraphs.
    """
    doc = Document()
    for raw in md_text.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:])
            p.style = "Heading 3"
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:])
            p.style = "Heading 2"
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:])
            p.style = "Heading 1"
        else:
            # strip basic code fences
            if line.startswith("```") or line.endswith("```"):
                continue
            doc.add_paragraph(line)
    doc.save(out_path)
    return out_path
